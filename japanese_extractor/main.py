import cv2
import os
from aip import AipOcr
from docx import Document
from docx.shared import Pt

image_folder = "picture"
if not os.path.exists(image_folder):
    os.makedirs(image_folder)


def get_file_content(file_path):
    with open(file_path, 'rb') as fp:
        return fp.read()


# OCR识别日语文本（在这里填写api信息）
def vcode2str(img_url):
    APP_ID = ""  # 填自己的信息获取的百度API信息
    API_KEY = ""
    SECRET_KEY = ""
    client = AipOcr(APP_ID, API_KEY, SECRET_KEY)

    # 获取图片的二进制数据
    image = get_file_content(img_url)

    # 设置OCR参数
    options = {}
    options["language_type"] = "JAP"
    options["detect_direction"] = "false"
    options["detect_language"] = "false"
    options["probability"] = "false"

    # 调用OCR API
    res = client.basicGeneral(image, options)

    # 提取文本
    strx = ""
    for tex in res.get("words_result", []):  # 遍历OCR结果
        strx += tex["words"] + "\n"  # 每一行文字加入到字符串，换行

    if not strx.strip():
        print(f"警告：OCR未提取到任何文字，图像可能不清晰或无文字：{img_url}")

    return strx


# 提取视频每两秒一帧并保存
def read_video(video_file):
    if not os.path.exists(video_file):
        print(f"错误：视频文件 {video_file} 不存在！")
        return

    cap = cv2.VideoCapture(video_file)

    if not cap.isOpened():
        print(f"错误：无法打开视频文件 {video_file}！")
        return

    # 获取视频的帧率和总帧数
    fps = cap.get(cv2.CAP_PROP_FPS)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))  # 获取视频的总帧数

    print(f"视频帧率：{fps} FPS, 总帧数：{total_frames}")  # 打印视频的帧率和总帧数

    frame_count = 0
    image_count = 0

    image_files = []

    while True:
        ret, frame = cap.read()

        if not ret:
            break  # 到达视频结尾

        # 每两秒提取一帧，基于 fps 计算
        if frame_count % int(2 * fps) == 0:  # 每 2 秒提取一次图像

            # 保存图像，并记录文件名
            image_name = f'{image_folder}/image{image_count + 1}.jpg'
            cv2.imwrite(image_name, frame)  # 保存原始图像
            print(f"已保存: {image_name}")

            image_files.append((image_count, image_name))  # 记录图像的顺序及文件名
            image_count += 1

        frame_count += 1

    cap.release()
    return image_files


def write_to_word(texts):
    doc = Document()

    # 设置黑体字
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(12)

    for i, text in enumerate(texts, start=1):
        doc.add_paragraph(f"{i}. {text}", style='Normal')
        doc.add_paragraph()

    doc.save("output1.docx")
    print(f"文本已成功写入到 output1.docx")


# 主流程
def process_video(video_file):

    print("提取视频帧...")
    image_files = read_video(video_file)

    print(f"提取到 {len(image_files)} 张图片进行OCR识别。")
    extracted_texts = [""] * len(image_files)

    for index, img_file in image_files:
        print(f"处理图片: {img_file}")
        text = vcode2str(img_file)
        extracted_texts[index] = text.strip()  # 根据图像顺序将文本存储到对应位置

    # 第三步：将提取的文本写入Word文档
    if any(extracted_texts):
        print("将提取的文本写入 Word 文档...")
        write_to_word(extracted_texts)
        print("日语文字已成功写入到 output1.docx")
    else:
        print("没有提取到任何文本，未生成 Word 文档。")


# 调用主流程
video_file = ""
process_video(video_file)
